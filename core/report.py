from docx import Document
from datetime import datetime
import pandas as pd
from pathlib import Path
import os

# Base directory du projet
BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = BASE_DIR / "data" / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def generate_audit_report(
    nom_client="",
    gap_analysis_file=OUTPUT_DIR / "gap_analysis.xlsx",
    output_file=None
):
    """
    Génère un rapport Word ISO 27001 à partir de la Gap Analysis déjà enregistrée.
    Le nom du client est automatiquement lu depuis l’Excel si dispo.
    """

    # Vérifier que la Gap Analysis existe
    if not Path(gap_analysis_file).exists():
        raise FileNotFoundError(f"❌ Fichier Gap Analysis introuvable : {gap_analysis_file}")

    # Charger la Gap Analysis depuis l’Excel
    df_gap = pd.read_excel(gap_analysis_file)

    # Si le nom du client est dans l’Excel, on le récupère
    if "Nom du client" in df_gap.columns and df_gap["Nom du client"].notna().any():
        nom_client = df_gap["Nom du client"].dropna().unique()[0]

    # Définir le nom de sortie si pas fourni
    if not output_file:
        safe_client_name = nom_client.replace(" ", "_") if nom_client else "Audit"
        output_file = OUTPUT_DIR / f"rapport_audit_{safe_client_name}.docx"

    # Créer le document Word
    doc = Document()
    date_document = datetime.now().strftime("%d/%m/%Y")

    # --- En-tête ---
    doc.add_heading(f"Rapport d'Audit ISO 27001 - {nom_client}", level=1)
    doc.add_paragraph(f"Date : {date_document}")
    doc.add_paragraph("")

    # --- Contenu : boucle sur la Gap Analysis ---
    for _, row in df_gap.iterrows():
        domaine = row.get('Domaine ISO 27001', '')
        clause = row.get('Clause', '')
        question = row.get('Question', '')
        statut = row.get('Statut', '')
        reponse = row.get('Réponse', '')
        justification = row.get('Justification', '')
        reco = row.get('Recommandation', '')
        question_comp = row.get('Question complémentaire', '')

        # Titre domaine
        if clause:
            doc.add_heading(f"{domaine} - Clause {clause}", level=2)
        else:
            doc.add_heading(f"{domaine}", level=2)

        # Détails
        doc.add_paragraph(f"Question : {question}")
        doc.add_paragraph(f"Statut : {statut}")
        doc.add_paragraph(f"Réponse : {reponse}")
        if justification:
            doc.add_paragraph(f"Justification : {justification}")
        doc.add_paragraph(f"Recommandation : {reco}")
        if question_comp:
            doc.add_paragraph(f"Question complémentaire : {question_comp}")
        doc.add_paragraph("")

    # Sauvegarder le document
    doc.save(output_file)
    print(f"✅ Rapport généré : {output_file}")
    return output_file
