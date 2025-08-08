from docx import Document
from datetime import datetime
import pandas as pd
import os

def generate_audit_report(
    nom_client="",
    gap_analysis_file="data/output/gap_analysis.xlsx",
    output_file=None
):
    """
    Génère un rapport Word d'audit ISO 27001 à partir de la Gap Analysis.
    Ajoute automatiquement le nom du client et la date du jour.
    """

    # Vérifier existence du fichier Excel
    if not os.path.exists(gap_analysis_file):
        raise FileNotFoundError(f"❌ Fichier Gap Analysis introuvable : {gap_analysis_file}")

    # Charger la Gap Analysis
    gap_df = pd.read_excel(gap_analysis_file)
    gap_analysis = gap_df.to_dict(orient="records")

    # Charger le template Word
    template_path = "templates/rapport_audit_template.docx"
    if os.path.exists(template_path):
        doc = Document(template_path)
    else:
        # Si pas de template → créer un document vide
        doc = Document()

    # Date du rapport
    date_document = datetime.now().strftime("%d/%m/%Y")

    # --- Remplir éventuellement les zones dans le template ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip().lower()
                if "version" in text:
                    try:
                        row.cells[1].text = date_document
                    except:
                        pass
                if "client" in text and nom_client:
                    try:
                        row.cells[1].text = nom_client
                    except:
                        pass

    # --- Titre principal ---
    doc.add_heading(f"Rapport d'Audit ISO 27001 - {nom_client}", level=1)
    doc.add_paragraph(f"Date : {date_document}")
    doc.add_paragraph("")

    # --- Ajout des résultats de la Gap Analysis ---
    for entry in gap_analysis:
        domaine = entry.get('Domaine', entry.get('Domaine ISO 27001', ''))
        clause = entry.get('Clause', '')
        question = entry.get('Question', '')
        statut = entry.get('Statut', '')
        reponse = entry.get('Réponse', '')
        justification = entry.get('Justification', '')
        reco = entry.get('Recommandation', '')
        question_comp = entry.get('Question complémentaire', '')

        # Titre domaine + clause
        if clause:
            doc.add_heading(f"{domaine} - Clause {clause}", level=2)
        else:
            doc.add_heading(f"{domaine}", level=2)

        # Détails
        doc.add_paragraph(f"Question : {question}")
        doc.add_paragraph(f"Statut : {statut}")
        doc.add_paragraph(f"Réponse : {reponse}")
        if justification:
            doc.add_paragraph(f"Justification : {justification}")
        doc.add_paragraph(f"Recommandation : {reco}")
        if question_comp:
            doc.add_paragraph(f"Question complémentaire : {question_comp}")
        doc.add_paragraph("")

    # --- Sauvegarder le rapport ---
    if not output_file:
        # Nom par défaut = client + date
        safe_client_name = nom_client.replace(" ", "_").replace("/", "-")
        output_file = f"data/output/rapport_audit_{safe_client_name}.docx"

    doc.save(output_file)
    print(f"✅ Rapport généré : {output_file}")
    return output_file
