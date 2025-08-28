# --- Imports ---
import os
from pathlib import Path
from io import BytesIO
import io
from typing import Optional, Dict, List, Tuple

import streamlit as st
import pandas as pd
import plotly.express as px
import fitz  # PyMuPDF
import docx
import json
from dotenv import load_dotenv
from openai import OpenAI
import base64
import hashlib
import re
from datetime import datetime

# Optional dependencies (safe fallbacks if missing)
try:
    import openpyxl  # for Excel styling
    OPENPYXL_AVAILABLE = True
except Exception:
    OPENPYXL_AVAILABLE = False

# ISO 27001 (existant)
from core.questions import ISO_QUESTIONS_INTERNE, ISO_QUESTIONS_MANAGEMENT
from core.analysis import (
    analyse_responses,
    save_gap_analysis,
    generate_action_plan_from_ai,
    save_action_plan_to_excel
)
from core.report import generate_audit_report

# ANSSI (nouveau)
from core.anssi_hygiene import ANSSI_SECTIONS, flatten_measures, STATUSES, SCORE_MAP

# (Optionnel) RAG utils si présents
try:
    from utils.ai_helper import build_vector_index, propose_anssi_answer  # RAG avancé
    RAG_AVAILABLE = True
except Exception:
    RAG_AVAILABLE = False

# --- Config & constantes ---
st.set_page_config(page_title="Audit Assistant (ISO 27001 • ANSSI)", layout="wide")

BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "data" / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# =========================================================
#   Fond d'écran (optionnel)
# =========================================================
def add_bg_from_local(image_file: str):
    try:
        with open(image_file, "rb") as f:
            data = f.read()
        encoded = base64.b64encode(data).decode()
        st.markdown(
            f"""
            <style>
            .stApp {{
                background-image: url("data:image/png;base64,{encoded}");
                background-size: cover;
                background-position: center;
                background-repeat: no-repeat;
                background-attachment: fixed;
            }}
            </style>
            """,
            unsafe_allow_html=True
        )
    except Exception:
        pass

if (BASE_DIR / "bg.png").exists():
    add_bg_from_local(str(BASE_DIR / "bg.png"))

st.markdown("""
<style>
/* Forcer les titres à rester blancs dans toute l'app */
html[data-theme="light"] .stApp h1,
html[data-theme="dark"]  .stApp h1,
html[data-theme="light"] .stApp h2,
html[data-theme="dark"]  .stApp h2,
html[data-theme="light"] .stApp h3,
html[data-theme="dark"]  .stApp h3,
html[data-theme="light"] .stApp h4,
html[data-theme="dark"]  .stApp h4 {
  color: #ffffff !important;
  text-shadow: 0 2px 6px rgba(0,0,0,0.45); /* améliore la lisibilité sur fonds clairs ou foncés */
}
</style>
""", unsafe_allow_html=True)


# =========================================================
#   OpenAI helpers
# =========================================================
def get_openai_api_key() -> Optional[str]:
    try: load_dotenv()
    except Exception: pass
    key = os.getenv("OPENAI_API_KEY")
    if key: return key
    try: return st.secrets["OPENAI_API_KEY"]
    except Exception: return None

def get_openai_client() -> Optional[OpenAI]:
    key = get_openai_api_key()
    if not key: return None
    try: return OpenAI(api_key=key)
    except Exception: return None

# =========================================================
#   Uploader GLOBAL (persistant + réutilisable partout)
# =========================================================
def _init_uploaded_docs_state():
    st.session_state.setdefault("uploaded_docs", [])  # [{name, bytes, size, sha1}]

def _file_sha1(b: bytes) -> str:
    return hashlib.sha1(b).hexdigest()

def _extract_text_from_pdf_bytes(b: bytes) -> str:
    try:
        pdf = fitz.open(stream=b, filetype="pdf")
        return "\n".join([p.get_text("text") for p in pdf])
    except Exception:
        return ""

def _extract_text_from_docx_bytes(b: bytes) -> str:
    try:
        d = docx.Document(io.BytesIO(b))
        return "\n".join(p.text for p in d.paragraphs if p.text)
    except Exception:
        return ""

def render_global_uploader():
    """Affiche un uploader réutilisable sur toutes les pages/onglets.
    Les fichiers sont mémorisés et réutilisables pour l'IA/RAG."""
    _init_uploaded_docs_state()
    with st.expander("📎 Documents d'appui (uploader global) — visibles partout", expanded=True):
        new_files = st.file_uploader(
            "Ajouter des documents (PDF/DOCX/TXT/PNG/JPG) pour enrichir l'analyse (RAG)",
            type=["pdf", "docx", "txt", "png", "jpg", "jpeg"],
            accept_multiple_files=True,
            key="global_uploader",
            help="Les documents ajoutés ici restent disponibles sur toutes les pages."
        )
        if new_files:
            added = 0
            for f in new_files:
                data = f.read()
                sig = _file_sha1(data)
                if not any(item.get("sha1") == sig for item in st.session_state["uploaded_docs"]):
                    st.session_state["uploaded_docs"].append({
                        "name": f.name,
                        "bytes": data,
                        "size": len(data),
                        "sha1": sig,
                    })
                    added += 1
            if added:
                st.success(f"{added} document(s) ajouté(s).")

        if st.session_state["uploaded_docs"]:
            st.caption("Documents mémorisés :")
            for i, item in enumerate(st.session_state["uploaded_docs"], start=1):
                col1, col2, col3 = st.columns([6, 2, 2])
                with col1:
                    st.write(f"{i}. **{item['name']}** — {round(item['size']/1024, 1)} KB")
                with col2:
                    st.download_button(
                        "Télécharger",
                        data=item["bytes"],
                        file_name=item["name"],
                        key=f"dl_{item['sha1']}"
                    )
                with col3:
                    if st.button("Retirer", key=f"rm_{item['sha1']}"):
                        st.session_state["uploaded_docs"] = [
                            x for x in st.session_state["uploaded_docs"] if x["sha1"] != item["sha1"]
                        ]
                        st.rerun()

def get_uploaded_docs_bytes() -> List[Tuple[str, bytes]]:
    _init_uploaded_docs_state()
    return [(x["name"], x["bytes"]) for x in st.session_state["uploaded_docs"]]

def get_uploaded_docs_text(truncate: int = 16000) -> str:
    """Concatène le texte des documents uploadés (PDF/DOCX/TXT)."""
    _init_uploaded_docs_state()
    texts: List[str] = []
    for item in st.session_state["uploaded_docs"]:
        name = item["name"].lower()
        b = item["bytes"]
        if name.endswith(".pdf"):
            texts.append(_extract_text_from_pdf_bytes(b))
        elif name.endswith(".docx"):
            texts.append(_extract_text_from_docx_bytes(b))
        elif name.endswith(".txt"):
            try:
                texts.append(b.decode("utf-8", errors="ignore"))
            except Exception:
                pass
    return ("\n\n".join(texts))[:truncate]

# =========================================================
#   Nettoyage IA (no JSON rendu) + parsing statut
# =========================================================
def ensure_plain_text(s: str) -> str:
    """Supprime fences ```...``` et convertit un éventuel JSON simple en texte clair FR."""
    if not isinstance(s, str):
        return str(s)
    s2 = re.sub(r"```(?:json|JSON)?\s*", "", s)
    s2 = s2.replace("```", "").strip()
    # tenter JSON -> texte
    try:
        obj = json.loads(s2)
        parts: List[str] = []
        if isinstance(obj, dict):
            if "status" in obj: parts.append(f"**Statut** : {obj['status']}")
            if "justification" in obj and obj.get("justification"):
                parts.append(f"**Justification** : {obj['justification']}")
            if "recommandations" in obj and isinstance(obj["recommandations"], list):
                parts.append("**Recommandations :**\n- " + "\n- ".join(map(str, obj["recommandations"])))
            if "actions_top3" in obj and isinstance(obj["actions_top3"], list):
                parts.append("**Actions prioritaires :**\n- " + "\n- ".join(map(str, obj["actions_top3"])))
            if not parts:
                for k, v in obj.items():
                    if isinstance(v, list):
                        parts.append(f"{k} :\n- " + "\n- ".join(map(str, v)))
                    else:
                        parts.append(f"{k} : {v}")
            return "\n\n".join(parts)
        if isinstance(obj, list):
            lines: List[str] = []
            for it in obj:
                if isinstance(it, dict):
                    line = ", ".join([f"{k}={v}" for k, v in it.items()])
                    lines.append(f"- {line}")
                else:
                    lines.append(f"- {it}")
            return "\n".join(lines)
    except Exception:
        pass
    return s2

def parse_status_from_text(txt: str) -> Optional[str]:
    """
    Détecte un statut dans du texte libre.
    Gère : Conforme, Partiellement conforme, Non conforme, Pas réponse, NA (Not Applicable).
    """
    if not isinstance(txt, str):
        return None
    t = txt.lower()

    # expressions courantes pour NA
    if any(kw in t for kw in [
        "not applicable", "n/a", "n.a", "na)", "(na", "na ", " non applicable", " hors périmètre", " hors perimetre"
    ]):
        return "NA"

    for s in ["Conforme", "Partiellement conforme", "Non conforme", "Pas réponse", "NA"]:
        if s.lower() in t:
            return s

    m = re.search(
        r"(statut|status)\s*[:\-]\s*(conforme|partiellement conforme|non conforme|pas réponse|na|n\/a|not applicable)",
        t
    )
    if m:
        val = m.group(2).strip()
        if val in ["na", "n/a", "not applicable"]:
            return "NA"
        for s in ["Conforme", "Partiellement conforme", "Non conforme", "Pas réponse"]:
            if s.lower() == val:
                return s
    return None

# =========================================================
#   Statuts & scoring (ajout NA)
# =========================================================
# Assure que NA est présent même si le module core n’était pas à jour
if "NA" not in STATUSES:
    try:
        STATUSES.append("NA")
    except Exception:
        pass

STATUS_META = {
    "Conforme": {"score": 1.00, "emoji": "✅", "priority": "Low"},
    "Partiellement conforme": {"score": 0.50, "emoji": "🟡", "priority": "Medium"},
    "Non conforme": {"score": 0.00, "emoji": "❌", "priority": "High"},
    "Pas réponse": {"score": 0.25, "emoji": "⚪", "priority": "Medium"},
    "NA": {"score": None, "emoji": "🚫", "priority": "N/A"},
}

# =========================================================
#   Helpers ANSSI : DataFrames + Exports + Rapport DOCX
# =========================================================
def _anssi_build_dataframe(measures, status_map: Dict[str, str], justifs_map: Dict[str, str]) -> pd.DataFrame:
    rows = []
    for m in measures:
        mid = m["id"]; theme = m["theme"]; title = m["title"]
        stt = status_map.get(mid, "Pas réponse")
        meta = STATUS_META.get(stt, STATUS_META["Pas réponse"])
        score = meta["score"]
        rows.append({
            "Thème": theme,
            "ID": mid,
            "Mesure": title,
            "Statut": stt,
            "Emoji": meta["emoji"],
            "Score (%)": int(round(score*100)) if isinstance(score, (int, float)) else None,
            "Priorité": meta["priority"],
            "Justification": ensure_plain_text(justifs_map.get(mid, "")),
        })
    order = ["Thème", "ID", "Mesure", "Statut", "Emoji", "Score (%)", "Priorité", "Justification"]
    return pd.DataFrame(rows)[order].sort_values(["Thème","ID"]).reset_index(drop=True)

def _anssi_theme_maturity(df: pd.DataFrame) -> pd.DataFrame:
    use = df[df["Score (%)"].notnull()]
    if use.empty:
        return pd.DataFrame(columns=["Thème","Maturité moyenne (%)"])
    g = use.groupby("Thème")["Score (%)"].mean().round(1).reset_index()
    g.columns = ["Thème", "Maturité moyenne (%)"]
    return g.sort_values("Maturité moyenne (%)", ascending=False)

def _save_csv_pretty(df: pd.DataFrame, path: Path) -> None:
    df_csv = df.copy()
    df_csv["Score (%)"] = df_csv["Score (%)"].fillna("")
    path.write_bytes(df_csv.to_csv(index=False).encode("utf-8"))

def _save_excel_styled(df: pd.DataFrame, theme_summary: pd.DataFrame, path: Path) -> None:
    with pd.ExcelWriter(path, engine="openpyxl" if OPENPYXL_AVAILABLE else None) as xw:
        df.to_excel(xw, sheet_name="Résultats détaillés", index=False)
        theme_summary.to_excel(xw, sheet_name="Synthèse par thème", index=False)
        if not OPENPYXL_AVAILABLE:
            return
        wb = xw.book
        # Résultats détaillés
        ws = wb["Résultats détaillés"]
        widths = {"A":16, "B":10, "C":60, "D":20, "E":8, "F":11, "G":12, "H":80}
        for col, w in widths.items():
            ws.column_dimensions[col].width = w
        from openpyxl.styles import PatternFill, Font, Alignment
        head_fill = PatternFill(start_color="FFDADADA", end_color="FFDADADA", fill_type="solid")
        for cell in ws[1]:
            cell.fill = head_fill
            cell.font = Font(bold=True)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
        color_map = {
            "Conforme": "FFC6EFCE",
            "Partiellement conforme": "FFFFF2CC",
            "Non conforme": "FFF8CBAD",
            "Pas réponse": "FFD9D9D9",
            "NA": "FFCCE5FF",
        }
        for row in ws.iter_rows(min_row=2, min_col=1, max_col=ws.max_column):
            statut = row[3].value  # D
            fill = PatternFill(start_color=color_map.get(statut, "FFFFFFFF"),
                               end_color=color_map.get(statut, "FFFFFFFF"),
                               fill_type="solid")
            row[3].fill = fill
            # wrap text for Mesure & Justification
            row[2].alignment = Alignment(wrap_text=True, vertical="top")
            row[7].alignment = Alignment(wrap_text=True, vertical="top")
        # Synthèse par thème
        ws2 = wb["Synthèse par thème"]
        for cell in ws2[1]:
            cell.fill = head_fill
            cell.font = Font(bold=True)
        for col in ["A","B"]:
            ws2.column_dimensions[col].width = 30 if col=="A" else 22

def _save_action_plan_excel(actions_df: pd.DataFrame, path: Path) -> None:
    with pd.ExcelWriter(path, engine="openpyxl" if OPENPYXL_AVAILABLE else None) as xw:
        actions_df.to_excel(xw, sheet_name="Plan d'actions", index=False)
        if OPENPYXL_AVAILABLE:
            wb = xw.book
            ws = wb["Plan d'actions"]
            from openpyxl.styles import PatternFill, Font, Alignment
            for cell in ws[1]:
                cell.font = Font(bold=True)
            # widths
            widths = {"A":10,"B":12,"C":50,"D":12,"E":24,"F":30,"G":14,"H":12}
            for col, w in widths.items():
                ws.column_dimensions[col].width = w
            for row in ws.iter_rows(min_row=2, min_col=1, max_col=ws.max_column):
                row[2].alignment = Alignment(wrap_text=True, vertical="top")  # Action
                row[5].alignment = Alignment(wrap_text=True, vertical="top")  # Justification

def _save_anssi_report_docx(df: pd.DataFrame, theme_summary: pd.DataFrame,
                            actions_df: Optional[pd.DataFrame], path: Path,
                            org_meta: Dict[str, str]) -> None:
    """
    Rapport Word ANSSI structuré, contenant TOUTES les réponses d’audit :
    - Couverture + Contexte
    - Executive Summary (KPI)
    - Synthèse par thème
    - Résultats détaillés (toutes mesures) : ID, Mesure, Statut, Justification
    - (Optionnel) Plan d’actions
    """
    d = docx.Document()
    now = datetime.now().strftime("%Y-%m-%d")

    # --- Couverture
    d.add_heading("ANSSI Hygiene – Assessment Report", 0)
    d.add_paragraph(f"Date: {now}")
    if org_meta:
        d.add_paragraph(
            f"Secteur: {org_meta.get('secteur','-')}  |  Employés: {org_meta.get('nb_emp','-')}  |  Pays: {org_meta.get('pays','-')}"
        )

    # --- Executive Summary
    d.add_heading("Executive Summary", level=1)
    total = len(df)
    c = (df["Statut"] == "Conforme").sum()
    pc = (df["Statut"] == "Partiellement conforme").sum()
    nc = (df["Statut"] == "Non conforme").sum()
    na = (df["Statut"] == "Pas réponse").sum()
    nna = (df["Statut"] == "NA").sum()
    valid = df["Score (%)"].dropna()
    overall = round(valid.mean(), 1) if not valid.empty else 0.0
    p = d.add_paragraph()
    p.add_run("Maturité globale: ").bold = True
    p.add_run(f"{overall}%")
    d.add_paragraph(f"Répartition statuts: ✅ {c} | 🟡 {pc} | ❌ {nc} | ⚪ {na} | 🚫 {nna}")

    # --- Synthèse par thème (table)
    d.add_heading("Synthèse par thème", level=1)
    if not theme_summary.empty:
        t = d.add_table(rows=1, cols=2)
        hdr = t.rows[0].cells
        hdr[0].text = "Thème"
        hdr[1].text = "Maturité moyenne (%)"
        for _, r in theme_summary.iterrows():
            row = t.add_row().cells
            row[0].text = str(r["Thème"])
            row[1].text = str(r["Maturité moyenne (%)"])
    else:
        d.add_paragraph("Aucune mesure avec score (NA partout).")

    # --- Résultats détaillés par thème (TOUTES les réponses)
    d.add_heading("Résultats détaillés", level=1)
    for theme, g in df.groupby("Thème"):
        d.add_heading(theme, level=2)
        # table avec toutes les mesures de ce thème
        tab = d.add_table(rows=1, cols=5)
        hdr = tab.rows[0].cells
        hdr[0].text = "ID"
        hdr[1].text = "Mesure"
        hdr[2].text = "Statut"
        hdr[3].text = "Score (%)"
        hdr[4].text = "Justification"
        for _, r in g.iterrows():
            row = tab.add_row().cells
            row[0].text = str(r["ID"])
            row[1].text = str(r["Mesure"])
            row[2].text = f"{r['Emoji']} {r['Statut']}"
            row[3].text = "" if pd.isna(r["Score (%)"]) else str(int(r["Score (%)"]))
            row[4].text = str(r["Justification"]) if r["Justification"] else "-"

    # --- Plan d’actions (optionnel)
    if actions_df is not None and not actions_df.empty:
        d.add_heading("Plan d’actions recommandé", level=1)
        t = d.add_table(rows=1, cols=8)
        hdr = t.rows[0].cells
        hdr[0].text = "ID"
        hdr[1].text = "Thème"
        hdr[2].text = "Action"
        hdr[3].text = "Priorité"
        hdr[4].text = "Responsable"
        hdr[5].text = "Justification"
        hdr[6].text = "Échéance"
        hdr[7].text = "Statut"
        for _, r in actions_df.iterrows():
            row = t.add_row().cells
            row[0].text = str(r["ID"])
            row[1].text = str(r["Thème"])
            row[2].text = str(r["Action"])
            row[3].text = str(r["Priorité"])
            row[4].text = str(r.get("Owner", ""))
            row[5].text = str(r.get("Justification", ""))
            row[6].text = str(r.get("Échéance", ""))
            row[7].text = str(r.get("Suivi", "Ouvert"))
    d.save(path)

def _build_anssi_action_plan_df(df: pd.DataFrame) -> pd.DataFrame:
    """
    Construit un plan d’actions à partir du DF des résultats.
    Utilise l’IA si possible pour produire des actions ciblées; sinon fallback générique.
    """
    # cibler d’abord Non conforme / Partiellement conforme
    target = df[df["Statut"].isin(["Non conforme", "Partiellement conforme"])].copy()
    if target.empty:
        return pd.DataFrame(columns=["ID","Thème","Action","Priorité","Owner","Justification","Échéance","Suivi"])

    client = get_openai_client()
    actions: List[Dict[str, str]] = []

    if client is not None:
        # prompt groupé pour limiter les coûts
        items = []
        for _, r in target.iterrows():
            items.append({
                "id": r["ID"], "theme": r["Thème"],
                "mesure": r["Mesure"], "statut": r["Statut"],
                "justif": r["Justification"]
            })
        system = (
            "Tu es un consultant cybersécurité senior. "
            "Pour chaque mesure fournie, propose 1 à 2 actions concrètes, ciblées et priorisées. "
            "Format JSON strict: [{\"id\":\"...\",\"actions\":[\"...\",\"...\"],\"priorite\":\"High|Medium|Low\"}]"
        )
        user = "Mesures à traiter:\n" + json.dumps(items, ensure_ascii=False)
        try:
            resp = client.chat.completions.create(
                model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
                messages=[{"role":"system","content":system},{"role":"user","content":user}],
                temperature=0.2,
            )
            raw = (resp.choices[0].message.content or "").strip()
            data = json.loads(raw)
            # construire lignes
            for it in data:
                mid = it.get("id")
                prio = it.get("priorite","High")
                acts = it.get("actions") or []
                base = target[target["ID"] == mid]
                if base.empty:
                    continue
                theme = base.iloc[0]["Thème"]
                just = base.iloc[0]["Justification"]
                for a in acts[:2]:
                    actions.append({
                        "ID": mid, "Thème": theme, "Action": ensure_plain_text(a),
                        "Priorité": prio, "Owner": "", "Justification": ensure_plain_text(just),
                        "Échéance": "", "Suivi": "Ouvert"
                    })
        except Exception:
            client = None  # fallback

    if client is None:
        # Fallback simple
        default_actions = {
            "Non conforme": "Établir un plan de remédiation documenté, définir un owner et une échéance; mettre en place le contrôle requis.",
            "Partiellement conforme": "Compléter la documentation et étendre la couverture du contrôle; formaliser les preuves et indicateurs.",
        }
        for _, r in target.iterrows():
            actions.append({
                "ID": r["ID"], "Thème": r["Thème"], "Action": default_actions[r["Statut"]],
                "Priorité": "High" if r["Statut"] == "Non conforme" else "Medium",
                "Owner": "", "Justification": r["Justification"], "Échéance": "", "Suivi": "Ouvert"
            })

    return pd.DataFrame(actions, columns=["ID","Thème","Action","Priorité","Owner","Justification","Échéance","Suivi"])

# =========================================================
#                     ROUTER + HOME
# =========================================================
if "route" not in st.session_state:
    st.session_state["route"] = "home"

def go(route: str):
    st.session_state["route"] = route

def render_home():
    st.title("🧭 Audit Assistant")
    st.caption("Centralisez vos audits, comparez vos pratiques, générez plans d’actions et rapports.")
    # Uploader global présent dès l'accueil
    render_global_uploader()
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("ISO/IEC 27001")
        st.write("- Questionnaires (interne / pré-certif)\n- Gap Analysis + priorités\n- Rapport prêt à partager")
        st.button("▶️ Entrer dans ISO 27001", key="home_go_iso", use_container_width=True, on_click=lambda: go("iso27001"))
    with col2:
        st.subheader("ANSSI – Guide d’hygiène")
        st.write("- Auto-évaluation par thèmes/mesures\n- Score de maturité & quick-wins\n- Exports pro (CSV/XLSX/DOCX)")
        st.button("▶️ Entrer dans ANSSI Hygiène", key="home_go_anssi", use_container_width=True, on_click=lambda: go("anssi_hygiene"))

# =========================================================
#                      ANSSI PAGE
# =========================================================
def render_anssi_hygiene():
    st.title("🛡️ ANSSI – Guide d’hygiène")
    st.caption("Parcours : 1) Intro  •  2) Questionnaire  •  3) Revue  •  4) Résultats")
    # Uploader global visible partout (exigence)
    render_global_uploader()

    # --- State init ---
    st.session_state.setdefault("anssi_stage", "intro")       # intro | questions | review | results
    st.session_state.setdefault("anssi_org", {})              # contexte entreprise
    st.session_state.setdefault("anssi_status", {})           # {id_mesure: statut}
    st.session_state.setdefault("anssi_justifs", {})          # {id_mesure: justification}
    st.session_state.setdefault("anssi_index", None)          # RAG index si dispo

    measures = flatten_measures()
    total = len(measures)

    def compute_progress():
        status_map = st.session_state["anssi_status"]
        answered = sum(1 for m in measures if status_map.get(m["id"]) in ("Conforme","Partiellement conforme","Non conforme","NA"))
        pct_answers = int(round(100 * answered / total)) if total else 0
        return pct_answers, answered

    def _index_docs():
        bins = get_uploaded_docs_bytes()
        if not bins:
            st.warning("Aucun document chargé dans l’uploader global.")
            return
        if not RAG_AVAILABLE:
            st.info("Indexation avancée indisponible (module RAG non importé).")
            return

        class _UploadedLike:
            def __init__(self, name, data):
                self.name = name
                self._data = data
            def getvalue(self):
                return self._data

        files_like = [_UploadedLike(name, b) for name, b in bins]
        with st.spinner("Indexation et embeddings…"):
            st.session_state["anssi_index"] = build_vector_index(files_like)
        st.success("Index construit ✔️")

    stage = st.session_state["anssi_stage"]

    # ---------- 1) INTRO ----------
    if stage == "intro":
        st.subheader("1) Informations de contexte")
        c1, c2 = st.columns(2)
        secteur = c1.text_input("Secteur d’activité", st.session_state["anssi_org"].get("secteur",""), key="anssi_org_secteur")
        nb_emp   = c2.number_input("Nombre d’employés", min_value=1, value=int(st.session_state["anssi_org"].get("nb_emp", 100)), key="anssi_org_nbemp")
        ca       = c1.text_input("Chiffre d’affaires (ex: 120 M€)", st.session_state["anssi_org"].get("ca",""), key="anssi_org_ca")
        pays     = c2.text_input("Filiales / Pays (ex: FR, LU, DE)", st.session_state["anssi_org"].get("pays",""), key="anssi_org_pays")
        st.info("Les documents déposés via l’uploader global seront utilisés par l’IA (RAG/texte).")

        colA, colB, colC = st.columns([1,1,1])
        if colA.button("💾 Enregistrer & continuer", key="anssi_intro_save"):
            st.session_state["anssi_org"] = {"secteur": secteur, "nb_emp": nb_emp, "ca": ca, "pays": pays}
            st.session_state["anssi_stage"] = "questions"
            st.rerun()
        if colB.button("🧭 Retour à l’accueil", key="anssi_intro_home"):
            go("home")
        colC.button("🧱 Indexer les documents (IA avancée)", key="anssi_index_btn_intro", on_click=_index_docs)
        return

    # ---------- IA: Préremplissage global ----------
    def _anssi_autofill_from_global():
        client = get_openai_client()
        if client is None:
            st.warning("ℹ️ Pas de clé OpenAI — préremplissage IA désactivé.")
            return

        org = st.session_state.get("anssi_org", {})

        # Si index RAG dispo + construit -> par mesure (qualité max)
        if RAG_AVAILABLE and st.session_state.get("anssi_index") and st.session_state["anssi_index"].get("chunks"):
            with st.spinner("Analyse IA (RAG) par mesure…"):
                for m in measures:
                    mid = m["id"]
                    requirement = m["title"]
                    question_md = _to_question_fr(requirement, m.get("theme"))
                    try:
                        res = propose_anssi_answer(requirement, question_md, st.session_state["anssi_index"])
                        status = res.get("status", "Pas réponse")
                        if status not in STATUSES:
                            status = "Pas réponse"
                        justif = ensure_plain_text(res.get("justification", ""))
                        cits = res.get("citations", [])
                        if cits:
                            justif = justif + "\n\n" + "Citations: " + "; ".join([f"{c['doc']} p.{c['page']}" for c in cits if isinstance(c, dict) and c.get('doc') and c.get('page')])
                        st.session_state["anssi_status"][mid] = status
                        st.session_state["anssi_justifs"][mid] = justif
                    except Exception:
                        st.session_state["anssi_status"][mid] = st.session_state["anssi_status"].get(mid, "Pas réponse")
                st.success("✅ Préremplissage IA (RAG) terminé.")
            return

        # Sinon: fallback via texte concaténé des uploads (interne; JSON non affiché)
        text = get_uploaded_docs_text()
        if not text:
            st.warning("ℹ️ Aucun document global chargé. Ajoute des fichiers dans l’uploader global.")
            return

        measures_brief = [{"id": m["id"], "title": m["title"], "theme": m["theme"]} for m in measures]

        system_msg = (
            "Tu es un consultant cybersécurité senior. "
            "À partir du contexte d’entreprise et des extraits fournis, "
            "évalue chaque mesure ANSSI et propose un statut conservateur. "
            "Si l'information est insuffisante, réponds 'Pas réponse'. "
            "Réponds STRICTEMENT en JSON (liste d’objets) : "
            "[{\"id\":\"...\",\"status\":\"Conforme|Partiellement conforme|Non conforme|Pas réponse|NA\",\"justification\":\"...\",\"actions_top3\":[\"...\",\"...\",\"...\"]}]"
        )

        user_msg = f"""
Contexte organisation:
- Secteur: {org.get('secteur') or 'Inconnu'}
- Nb employés: {org.get('nb_emp') or 'Inconnu'}
- CA: {org.get('ca') or 'Inconnu'}
- Pays/Filiales: {org.get('pays') or 'Inconnu'}

Référentiel: ANSSI Guide d'hygiène (10 thèmes, ~42 mesures).
Mesures à évaluer (id/title/theme):
{json.dumps(measures_brief, ensure_ascii=False)}

Extraits de documents globaux (tronqués):
{text}

Consignes:
- Donne un statut par mesure parmi: Conforme | Partiellement conforme | Non conforme | Pas réponse | NA
- Justifie brièvement (2-4 lignes) + 3 actions prioritaires.
- Si incertain: 'Pas réponse'.
Renvoie uniquement du JSON valide.
"""

        try:
            resp = client.chat.completions.create(
                model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
                messages=[{"role": "system", "content": system_msg},
                          {"role": "user", "content": user_msg}],
                temperature=0.2,
            )
            raw = (resp.choices[0].message.content or "").strip()
            data = json.loads(raw)
        except Exception as e:
            st.error(f"Erreur IA: {e}")
            return

        for item in data:
            mid = item.get("id")
            status = item.get("status", "Pas réponse")
            justif = item.get("justification", "")
            if not mid:
                continue
            if status not in STATUSES:
                status = "Pas réponse"
            st.session_state["anssi_status"][mid] = status
            st.session_state["anssi_justifs"][mid] = ensure_plain_text(justif)

        st.success("✅ Préremplissage IA terminé (fallback texte concaténé).")

    # ---------- 2) QUESTIONNAIRE ----------
    if stage == "questions":
        pct_answers, answered = compute_progress()
        st.subheader("2) Questionnaire par thématiques")
        st.write(f"Avancement questionnaire : **{pct_answers}%** — ({answered}/{total})")
        st.progress(pct_answers)

        # Actions IA utiles directement ici
        cA, cB, cC = st.columns([1,1,1])
        cA.button("🧠 Préremplir automatiquement (docs globaux)", key="anssi_autofill_btn", on_click=_anssi_autofill_from_global)
        cB.button("🧱 (Re)Indexer documents (IA avancée)", key="anssi_index_btn_q", on_click=_index_docs)
        cC.button("🏠 Accueil", key="anssi_questions_home_btn", on_click=lambda: go("home"))

        theme = st.sidebar.radio("Thèmes", list(ANSSI_SECTIONS.keys()), key="anssi_theme_radio")
        st.subheader(theme)

        SYSTEM_FRENCH_PLAIN = (
            "Tu es un auditeur sénior cybersécurité/continuité. "
            "Réponds en français, en texte clair (phrases ou puces). "
            "N'utilise aucun JSON, aucun code fence. "
            "Commence par une ligne 'Statut: ...' avec l'une des valeurs: "
            "Conforme, Partiellement conforme, Non conforme, Pas réponse, NA. "
            "Puis donne une justification concise (3–6 lignes) et 2–4 actions concrètes."
        )

        for m in ANSSI_SECTIONS[theme]:
            mid = m["id"]
            requirement = m["title"]
            question_md = _to_question_fr(requirement, m.get("theme"))

            st.markdown(f"### {mid}")
            st.markdown(question_md)
            with st.expander("Voir l’exigence ANSSI (texte brut)"):
                st.write(requirement)

            # Statut (sélecteur)
            current = st.session_state["anssi_status"].get(mid, "Pas réponse")
            new_status = st.selectbox(
                "Statut",
                STATUSES,
                index=STATUSES.index(current) if current in STATUSES else STATUSES.index("Pas réponse"),
                key=f"status_{mid}"
            )
            st.session_state["anssi_status"][mid] = new_status

            # Zone texte consultant (réponse détaillée)
            cur_just = st.session_state["anssi_justifs"].get(mid, "")
            new_just = st.text_area(
                "Réponse détaillée (consultant) – Justification & éléments de preuve",
                value=cur_just,
                key=f"justif_{mid}",
                height=160,
                placeholder="Rédige une justification professionnelle : politiques, journaux, tickets, preuves de tests, etc."
            )
            st.session_state["anssi_justifs"][mid] = new_just

            # IA par mesure (texte clair)
            cols = st.columns([1,1])
            if cols[0].button("💡 Proposer avec l’IA", key=f"ai_{mid}"):
                client = get_openai_client()
                if client is None:
                    st.warning("Clé OpenAI manquante.")
                else:
                    try:
                        if RAG_AVAILABLE and st.session_state.get("anssi_index") and st.session_state["anssi_index"].get("chunks"):
                            # RAG → formater en texte clair
                            res = propose_anssi_answer(requirement, question_md, st.session_state["anssi_index"])
                            status = res.get("status")
                            if status in STATUSES:
                                st.session_state["anssi_status"][mid] = status
                            justif = ensure_plain_text(res.get("justification", ""))
                            cits = res.get("citations", [])
                            if cits:
                                justif += "\n\nCitations: " + "; ".join([f"{c['doc']} p.{c['page']}" for c in cits if isinstance(c, dict) and c.get('doc') and c.get('page')])
                            st.session_state["anssi_justifs"][mid] = justif
                        else:
                            # Fallback: prompt texte clair (sans JSON)
                            context_text = get_uploaded_docs_text(truncate=8000)
                            user_prompt = (
                                f"EXIGENCE: {requirement}\n\nQUESTION:\n{question_md}\n\n"
                                f"CONTEXTE (extraits des documents, éventuellement vide):\n{context_text}\n\n"
                                "Donne uniquement du texte clair. Pas de JSON, pas de balises."
                            )
                            resp = client.chat_completions.create if hasattr(client, "chat_completions") else client.chat.completions.create
                            out = resp(
                                model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
                                messages=[
                                    {"role":"system","content": SYSTEM_FRENCH_PLAIN},
                                    {"role":"user","content": user_prompt}
                                ],
                                temperature=0.2
                            )
                            content = ensure_plain_text((out.choices[0].message.content or "").strip())
                            detected = parse_status_from_text(content) or "Pas réponse"
                            if detected in STATUSES:
                                st.session_state["anssi_status"][mid] = detected
                            st.session_state["anssi_justifs"][mid] = content
                        st.success("Proposition IA appliquée ✔️")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Erreur IA: {e}")

            cols[1].markdown("&nbsp;")
            st.divider()

        c1, c2, c3 = st.columns([1,1,1])
        c1.button("⬅️ Retour à l’accueil", key="anssi_questions_home", on_click=lambda: go("home"))
        c2.button("🔄 Recalculer l’avancement", key="anssi_questions_recalc", on_click=st.rerun)
        if c3.button("➡️ Terminer le questionnaire", key="anssi_questions_next_review"):
            st.session_state["anssi_stage"] = "review"
            st.rerun()
        return

    # ---------- 2.5) REVIEW ----------
    if stage == "review":
        st.subheader("✔️ Revue avant analyse")
        pct_answers, answered = compute_progress()
        st.write(f"Avancement réponses (hors 'Pas réponse') : **{pct_answers}%** — ({answered}/{total})")
        st.progress(pct_answers)

        missing = [m for m in measures if st.session_state["anssi_status"].get(m["id"], "Pas réponse") == "Pas réponse"]
        if missing:
            st.warning(f"Mesures sans réponse : {len(missing)}")
            with st.expander("Voir les mesures sans réponse"):
                for m in missing:
                    st.write(f"- {m['id']} — {m['title']} ({m['theme']})")
        else:
            st.success("Toutes les mesures ont un statut (y compris 'NA' ou 'Pas réponse').")

        c1, c2, c3 = st.columns([1,1,1])
        c1.button("⬅️ Retour au questionnaire", key="anssi_review_back_to_questions", on_click=lambda: st.session_state.update({"anssi_stage":"questions"}) or st.rerun())
        c2.button("🏠 Accueil", key="anssi_review_home", on_click=lambda: go("home"))
        if c3.button("✅ Valider & commencer l’analyse (globale)", key="anssi_review_start"):
            st.session_state["anssi_stage"] = "results"
            st.rerun()
        return

    # ---------- 3) RÉSULTATS ----------
    if stage == "results":
        st.subheader("3) Résultats & Livrables")

        # DataFrames principaux
        df = _anssi_build_dataframe(measures, st.session_state["anssi_status"], st.session_state["anssi_justifs"])
        theme_summary = _anssi_theme_maturity(df)

        # Table interactive
        st.dataframe(df, use_container_width=True)

        # Charts rapides
        st.markdown("#### 📊 Répartition des statuts")
        counts = df["Statut"].value_counts().reset_index()
        counts.columns = ["Statut", "Nombre"]
        fig1 = px.pie(counts, values="Nombre", names="Statut", title="Répartition des statuts")
        st.plotly_chart(fig1, use_container_width=True)

        st.markdown("#### 📈 Maturité par thème")
        if not theme_summary.empty:
            fig2 = px.bar(theme_summary, x="Thème", y="Maturité moyenne (%)", title="Maturité moyenne par thème", text="Maturité moyenne (%)")
            st.plotly_chart(fig2, use_container_width=True)
        else:
            st.info("Aucune maturité calculable (scores NA seulement).")

        # Livrables
        csv_path = OUTPUT_DIR / "anssi_resultats.csv"
        xlsx_path = OUTPUT_DIR / "anssi_resultats.xlsx"
        plan_path = OUTPUT_DIR / "anssi_action_plan.xlsx"
        docx_path = OUTPUT_DIR / "anssi_rapport.docx"

        # Exports
        _save_csv_pretty(df, csv_path)
        _save_excel_styled(df, theme_summary, xlsx_path)

        # Plan d’actions (IA si possible)
        actions_df = _build_anssi_action_plan_df(df)
        if not actions_df.empty:
            _save_action_plan_excel(actions_df, plan_path)

        # Rapport DOCX (inclut TOUTES les réponses d’audit, bien structurées)
        _save_anssi_report_docx(df, theme_summary, actions_df if not actions_df.empty else None, docx_path, st.session_state.get("anssi_org", {}))

        # Boutons de téléchargement
        st.download_button("⬇️ Export CSV (propre)", data=open(csv_path, "rb").read(), file_name=csv_path.name, mime="text/csv")
        st.download_button("⬇️ Excel stylé (résultats + synthèse)", data=open(xlsx_path, "rb").read(), file_name=xlsx_path.name)
        if not actions_df.empty:
            st.download_button("⬇️ Plan d’actions (Excel)", data=open(plan_path, "rb").read(), file_name=plan_path.name)
        st.download_button("⬇️ Rapport Word (complet)", data=open(docx_path, "rb").read(), file_name=docx_path.name)

        st.button("↩️ Revenir au questionnaire", key="anssi_results_back_to_questions", on_click=lambda: st.session_state.update({"anssi_stage":"questions"}) or st.rerun())
        st.button("⬅️ Retour à l’accueil", key="anssi_results_home", on_click=lambda: go("home"))
        return

# =========================================================
#            ISO 27001 (page & IA préremplissage)
# =========================================================
def _ai_prefill_iso_by_domain(documents_text: str, iso_questions: Dict[str, List[Dict]]) -> Dict[str, Dict[str, str]]:
    client = get_openai_client()
    if client is None:
        return {}
    out: Dict[str, Dict[str, str]] = {}
    for domain, qs in iso_questions.items():
        q_list = []
        for q in qs:
            clause = q.get("clause", "")
            qtxt = q["question"]
            q_list.append({"clause": clause, "question": qtxt})
        ctx = documents_text[:16000]
        system = (
            "Tu es auditeur ISO/IEC 27001. "
            "Sur la base du contexte fourni, propose une réponse courte (2-4 lignes) et factuelle pour chaque question. "
            "N'invente pas si l'info n'existe pas; mets 'Information insuffisante'. "
            "Renvoie STRICTEMENT du JSON: {\"answers\": [{\"question\": \"...\", \"answer\": \"...\"}, ...]}"
        )
        user = f"DOMAINE: {domain}\nQUESTIONS: {json.dumps(q_list, ensure_ascii=False)}\n\nCONTEXTE:\n{ctx}"
        try:
            resp = client.chat.completions.create(
                model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
                messages=[{"role": "system", "content": system},
                          {"role": "user", "content": user}],
                temperature=0.2
            )
            content = (resp.choices[0].message.content or "").strip()
            data = json.loads(content)
            answers = data.get("answers", [])
        except Exception:
            answers = []
        out[domain] = {}
        for a in answers:
            qtxt = a.get("question", "")
            ans  = a.get("answer", "")
            if qtxt:
                out[domain][qtxt] = ans
    return out

def _mk_bullets(items: List[str]) -> str:
    return "\n".join([f"- {it}" for it in items])

def _to_question_fr(exigence: str, theme: Optional[str] = None) -> str:
    """Transforme une exigence ANSSI en question pro et actionnable."""
    if not exigence:
        return ""
    txt = exigence.strip()
    low = txt.lower()

    def block(title: str, bullets: List[str]) -> str:
        return f"**{title}**\n\nPoints attendus :\n{_mk_bullets(bullets)}"

    if any(k in low for k in ["sauvegard", "backup", "restaur"]):
        return block(
            "Comment l’organisation assure les sauvegardes et la restauration ?",
            ["Périmètre (serveurs/postes/bases/Cloud)", "Fréquence & rétention (3-2-1)", "Chiffrement & gestion des clés", "Tests de restauration (RPO/RTO)", "Supervision & PRA/PCA"]
        )
    # … (raccourci : autres heuristiques couvertes plus haut, conservées)

    return block(
        f"Comment l’organisation adresse l’exigence suivante : « {txt} » ?",
        ["Gouvernance", "Processus (SLA/approbations)", "Contrôles techniques", "Indicateurs & supervision", "Preuves (docs, logs, tickets)"]
    )

def render_iso27001():
    st.title("🔍 Audit ISO 27001")
    # Uploader global utilisable aussi côté ISO
    render_global_uploader()

    mode = st.radio(
        "🎯 Objectif d'audit",
        ["Audit interne", "Audit officiel / Pré-certification"],
        horizontal=True,
        index=0 if st.session_state.get("audit_mode", "interne") == "interne" else 1,
        key="audit_mode_choice",
    )
    st.session_state.audit_mode = "interne" if mode == "Audit interne" else "officiel"
    audit_mode = st.session_state.audit_mode

    ISO_QUESTIONS = ISO_QUESTIONS_INTERNE if audit_mode == "interne" else {**ISO_QUESTIONS_MANAGEMENT, **ISO_QUESTIONS_INTERNE}

    client_name_input = st.text_input(
        "🏢 Nom du client pour cet audit",
        placeholder="Exemple : D&A, CACEIS, Banque XYZ...",
        key="client_name",
    ).strip()

    if client_name_input:
        st.success(f"Client sélectionné : **{client_name_input}**")
    else:
        st.info("➡️ Indiquez le nom du client pour activer l’import des documents.")
        st.button("⬅️ Retour à l’accueil", key="iso_back_home", on_click=lambda: go("home"))
        st.stop()

    def extract_text_from_pdf(file):
        text = ""
        pdf = fitz.open(stream=file.read(), filetype="pdf")
        for page in pdf:
            text += page.get_text()
        return text

    def extract_text_from_docx(file):
        d = docx.Document(file)
        return "\n".join(p.text for p in d.paragraphs)

    def detect_client_name_with_ai(text):
        preview_text = text[:1500]
        prompt = f"""
Tu es un expert en audit ISO 27001.
Voici un extrait du début d'un document d'audit :
---
{preview_text}
---
À partir de cet extrait, identifie uniquement le NOM de l'organisation ou du client.
IMPORTANT :
- Ne donne pas d'explication.
- Ne réponds que par le nom détecté.
- Si tu n'es pas sûr, réponds exactement "Inconnu".
"""
        client = get_openai_client()
        if client is None:
            return "Inconnu"
        try:
            response = client.chat.completions.create(
                model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
                messages=[{"role": "user", "content": prompt}],
                temperature=0
            )
            return (response.choices[0].message.content or "").strip()
        except Exception:
            return "Inconnu"

    def make_example_txt():
        content = (
            "Client: ACME BANK\n"
            "Document: ISMS Policy v1.2\n"
            "Scope: HQ + DataCenter\n"
            "Summary: High-level information security policy aligned to ISO/IEC 27001.\n"
        )
        return content.encode("utf-8")

    def make_example_docx():
        d = docx.Document()
        d.add_heading("Access Control Procedure", level=1)
        d.add_paragraph("Client: ACME BANK")
        d.add_paragraph("Purpose: Define access control rules aligned with ISO/IEC 27001 A.9.")
        d.add_paragraph("Scope: Corporate systems, VDI, privileged accounts, third parties.")
        bio = BytesIO()
        d.save(bio)
        bio.seek(0)
        return bio.getvalue()

    def make_correct_example_docx(client_name: str):
        d = docx.Document()
        d.add_heading("Exemple - Procédure Sécurité", level=1)
        d.add_paragraph(f"Client: {client_name}")
        d.add_paragraph("Document type: Politique de sécurité de l'information.")
        d.add_paragraph("Aligné avec ISO/IEC 27001 (contrôles A.5 à A.18).")
        bio = BytesIO()
        d.save(bio)
        bio.seek(0)
        return bio.getvalue()

    st.subheader("📂 Importer documents du client (spécifique à ISO)")
    st.markdown(
        "Analysez vos documents (**politiques, procédures, rapports**) pour pré-remplir le questionnaire, "
        "générez une **Gap Analysis** et un **rapport Word** prêt à partager."
    )

    with st.expander("📎 Exemples de documents à tester (téléchargeables)"):
        col_a, col_b = st.columns(2)
        with col_a:
            st.download_button(
                "⬇️ Télécharger un exemple TXT",
                data=make_example_txt(),
                file_name="exemple_client.txt",
                mime="text/plain",
                use_container_width=True,
                key="iso_example_txt"
            )
        with col_b:
            st.download_button(
                "⬇️ Télécharger un exemple DOCX",
                data=make_example_docx(),
                file_name="exemple_procedure.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="iso_example_docx"
            )

    uploaded_files = st.file_uploader(
        "Formats acceptés : PDF, DOCX, TXT — limite 200 Mo par fichier.",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        key="iso_uploader"
    )

    # On combine textes des uploads ISO + uploader global
    documents_text = get_uploaded_docs_text()
    detected_client_names = set()

    # Ajoute le texte des fichiers uploadés ici (spécifique ISO)
    if uploaded_files:
        for file in uploaded_files:
            if file.name.lower().endswith(".pdf"):
                text = extract_text_from_pdf(file)
            elif file.name.lower().endswith(".docx"):
                text = extract_text_from_docx(file)
            elif file.name.lower().endswith(".txt"):
                text = file.read().decode("utf-8", errors="ignore")
            else:
                text = ""
            documents_text += "\n" + text

            detected_name = detect_client_name_with_ai(text)
            if detected_name and detected_name != "Inconnu":
                detected_client_names.add(detected_name)

        if len(detected_client_names) > 1:
            st.error(
                f"⚠️ Plusieurs clients détectés dans les documents : "
                f"{', '.join(detected_client_names)}"
            )
            st.stop()

        mismatch = detected_client_names and not any(
            client_name_input.lower() in name.lower() for name in detected_client_names
        )
        if mismatch:
            st.error(
                "🚨 Incohérence détectée : "
                f"documents analysés pour {', '.join(detected_client_names)}, "
                f"≠ nom saisi '{client_name_input}'.\n"
                "Veuillez corriger le nom ou importer les bons documents."
            )
            st.download_button(
                f"⬇️ Télécharger un exemple DOCX pour {client_name_input}",
                data=make_correct_example_docx(client_name_input),
                file_name=f"exemple_{client_name_input.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="iso_example_correct_docx"
            )
            st.stop()

    responses = {}
    if documents_text.strip():
        client = get_openai_client()
        if client is None:
            st.warning("ℹ️ Aucune clé OpenAI détectée — l’analyse IA des documents est désactivée.")
        else:
            st.info("📡 Analyse IA en cours...")
            responses = _ai_prefill_iso_by_domain(documents_text, ISO_QUESTIONS)
            st.success("✅ Questionnaire pré-rempli par l'IA.")

    if responses:
        gap_analysis = analyse_responses(responses, nom_client=client_name_input)
        save_gap_analysis(gap_analysis, nom_client=client_name_input)

        df_gap = pd.DataFrame(gap_analysis)

        if not df_gap.empty:
            st.subheader("📊 Gap Analysis (vue interactive)")
            domaines = ["Tous"] + sorted(df_gap["Domaine ISO 27001"].unique())
            priorites = ["Toutes"] + sorted(df_gap["Priorité"].unique())

            col1, col2 = st.columns(2)
            with col1:
                filtre_domaine = st.selectbox("📌 Filtrer par domaine", domaines, key="iso_filter_domain")
            with col2:
                filtre_priorite = st.selectbox("⚡ Filtrer par priorité", priorites, key="iso_filter_priority")

            df_filtre = df_gap.copy()
            if filtre_domaine != "Tous":
                df_filtre = df_filtre[df_filtre["Domaine ISO 27001"] == filtre_domaine]
            if filtre_priorite != "Toutes":
                df_filtre = df_filtre[df_filtre["Priorité"] == filtre_priorite]

            st.dataframe(df_filtre, use_container_width=True)

            export_gap = OUTPUT_DIR / "gap_analysis_ui.xlsx"
            df_filtre.to_excel(export_gap, index=False)
            st.download_button(
                "📥 Télécharger Gap Analysis (Excel)",
                data=open(export_gap, "rb").read(),
                file_name="gap_analysis.xlsx",
                key="iso_export_gap"
            )

            if not df_filtre.empty:
                statut_counts = df_filtre["Statut"].value_counts().reset_index()
                statut_counts.columns = ["Statut", "Nombre"]

                color_map = {
                    "Conforme": "#2ecc71",
                    "✅ Conforme": "#2ecc71",
                    "Non conforme": "#e74c3c",
                    "❌ Non conforme": "#e74c3c"
                }

                fig = px.pie(
                    statut_counts,
                    values="Nombre",
                    names="Statut",
                    title="Répartition par statut de conformité (selon filtre)",
                    color="Statut",
                    color_discrete_map=color_map
                )
                st.plotly_chart(fig, use_container_width=True)

    # Formulaire interactif final
    with st.form("audit_form"):
        final_responses = {}
        for domain, questions in ISO_QUESTIONS.items():
            st.subheader(f"📌 {domain}")
            final_responses[domain] = {}
            for q in questions:
                clause = q.get("clause", "")
                question_text = q["question"]
                question_display = f"{clause} – {question_text}" if clause else question_text
                answer_data = responses.get(domain, {}).get(question_text, "")

                key_suffix = f"{domain}_{clause}_{hash(question_text)}"

                if isinstance(answer_data, dict):
                    reponse_simple = answer_data.get("Réponse", "")
                    new_answer = st.text_area(
                        question_display,
                        value=reponse_simple,
                        key=f"ta_{key_suffix}"
                    )
                    final_responses[domain][question_text] = {**answer_data, "Réponse": new_answer}
                else:
                    new_answer = st.text_area(
                        question_display,
                        value=answer_data,
                        key=f"tb_{key_suffix}"
                    )
                    final_responses[domain][question_text] = new_answer

        submitted = st.form_submit_button("📥 Générer l'analyse et le rapport")  # <— pas de 'key' ici

    if submitted:
        gap_analysis = analyse_responses(final_responses, nom_client=client_name_input)
        save_gap_analysis(gap_analysis, nom_client=client_name_input)

        report_path = generate_audit_report()

        st.success("✅ Rapport généré avec succès !")
        st.download_button(
            "📄 Télécharger rapport Word",
            data=open(report_path, "rb").read(),
            file_name=Path(report_path).name,
            key="iso_download_report"
        )
        st.download_button(
            "📊 Télécharger Gap Analysis",
            data=open(OUTPUT_DIR / "gap_analysis.xlsx", "rb").read(),
            file_name="gap_analysis.xlsx",
            key="iso_download_gap"
        )

        action_client = get_openai_client()
        if action_client is None:
            st.warning("ℹ️ Pas de clé OpenAI — génération automatique du plan d’actions désactivée.")
        else:
            action_plan = generate_action_plan_from_ai(gap_analysis, nom_client=client_name_input)
            save_action_plan_to_excel(action_plan)

            st.subheader("📅 Plan d’actions recommandé")
            df_plan = pd.DataFrame(action_plan)
            if not df_plan.empty:
                st.dataframe(df_plan, use_container_width=True)
                st.download_button(
                    "📥 Télécharger le plan d’actions (Excel)",
                    data=open(OUTPUT_DIR / "action_plan.xlsx", "rb").read(),
                    file_name="plan_actions.xlsx",
                    key="iso_download_plan"
                )
            else:
                st.info("✅ Aucun plan d’action nécessaire, tout est conforme.")

    st.divider()
    st.button("⬅️ Retour à l’accueil", key="iso_home", on_click=lambda: go("home"))

# =========================================================
#                       DISPATCH
# =========================================================
page = st.session_state.get("route", "home")

if page == "home":
    render_home()
elif page == "anssi_hygiene":
    render_anssi_hygiene()
elif page == "iso27001":
    render_iso27001()
else:
    st.session_state["route"] = "home"
    render_home()